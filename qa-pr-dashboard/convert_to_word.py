#!/usr/bin/env python3
"""
Convert ARCHITECTURE.md to Word document format
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.style.font.size = Pt(18)
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.style.font.size = Pt(14)
            i += 1
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            heading.style.font.size = Pt(12)
            i += 1
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            i += 1
        # Handle code blocks
        elif line.startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            
            # Add code block as monospace paragraph
            if code_lines:
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'No Spacing'
                run = para.runs[0] if para.runs else para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
        # Handle empty lines
        elif line.strip() == '':
            doc.add_paragraph()
            i += 1
        # Handle regular text
        else:
            # Check if it's part of a code block (indented)
            if line.startswith('    ') or line.startswith('\t'):
                # Treat as code
                para = doc.add_paragraph(line.strip())
                run = para.runs[0] if para.runs else para.add_run(line.strip())
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                para.paragraph_format.left_indent = Inches(0.5)
            else:
                # Regular paragraph
                para = doc.add_paragraph(line.strip())
                para.paragraph_format.space_after = Pt(6)
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    md_file = Path(__file__).parent / 'ARCHITECTURE.md'
    docx_file = Path(__file__).parent / 'ARCHITECTURE.docx'
    
    if not md_file.exists():
        print(f"Error: {md_file} not found")
        sys.exit(1)
    
    parse_markdown_to_docx(md_file, docx_file)



